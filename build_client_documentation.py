from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables" / "ADONAK_Electronics_Client_Project_Documentation.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
MID_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E9F7F0"
LIGHT_GREY = "F5F7FA"
TEXT = "1F2937"
MUTED = "5B6B82"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in kwargs[edge]:
                    element.set(qn("w:{}".format(key)), str(kwargs[edge][key]))


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for i, width in enumerate(widths):
        grid.gridCol_lst[i].set(qn("w:w"), str(width))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def text_run(paragraph, text, size=10.5, bold=False, color=TEXT, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    text_run(p, text, 10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    text_run(p, text, 10.5)
    return p


def add_note(doc, lead, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9120])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GREEN)
    set_cell_margins(cell, 120, 160, 120, 160)
    set_cell_border(cell, top={"val": "single", "sz": 8, "color": "96D9BC"}, bottom={"val": "single", "sz": 8, "color": "96D9BC"}, left={"val": "single", "sz": 8, "color": "96D9BC"}, right={"val": "single", "sz": 8, "color": "96D9BC"})
    p = cell.paragraphs[0]
    text_run(p, lead + " ", 10.5, bold=True, color=MID_BLUE)
    text_run(p, body, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, title in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell)
        set_cell_border(cell, top={"val": "single", "sz": 4, "color": "B8C8DA"}, bottom={"val": "single", "sz": 4, "color": "B8C8DA"}, left={"val": "single", "sz": 4, "color": "B8C8DA"}, right={"val": "single", "sz": 4, "color": "B8C8DA"})
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        text_run(p, title, 9.5, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell, top={"val": "single", "sz": 3, "color": "D7E0EA"}, bottom={"val": "single", "sz": 3, "color": "D7E0EA"}, left={"val": "single", "sz": 3, "color": "D7E0EA"}, right={"val": "single", "sz": 3, "color": "D7E0EA"})
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(value, tuple):
                text_run(p, value[0], 10, bold=True, color=MID_BLUE)
                text_run(p, value[1], 10)
            else:
                text_run(p, str(value), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, MID_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_header_footer(section):
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text_run(p, "ADONAK ELECTRONICS SHOP", 8.5, bold=True, color=BLUE)
    text_run(p, "  |  Client Project Documentation", 8.5, color=MUTED)
    p.paragraph_format.space_after = Pt(0)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_run(p, "ADONAK Electronics Shop  |  Client Handover Guide  |  Page ", 8, color=MUTED)
    add_page_field(p)


def heading(doc, value, level=1):
    return doc.add_heading(value, level=level)


def para(doc, value, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and value.startswith(bold_lead):
        text_run(p, bold_lead, 11, bold=True, color=NAVY)
        text_run(p, value[len(bold_lead):], 11)
    else:
        text_run(p, value, 11)
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc.sections[0])

    # Cover page
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(8)
    text_run(p, "CLIENT HANDOVER GUIDE", 10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    text_run(p, "ADONAK", 29, bold=True, color=NAVY)
    text_run(p, " Electronics Shop", 29, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    text_run(p, "Project Documentation & Operating Guide", 16, color=MID_BLUE)

    cover = doc.add_table(rows=3, cols=2)
    set_table_geometry(cover, [2500, 6620])
    labels = ["Prepared for", "Document version", "Date"]
    values = ["ADONAK Electronics", "1.0 — Client Handover", "13 August 2026"]
    for r in range(3):
        for c in range(2):
            cell = cover.cell(r, c)
            set_cell_margins(cell, 110, 150, 110, 150)
            set_cell_border(cell, top={"val": "single", "sz": 5, "color": "C9D6E4"}, bottom={"val": "single", "sz": 5, "color": "C9D6E4"}, left={"val": "single", "sz": 5, "color": "C9D6E4"}, right={"val": "single", "sz": 5, "color": "C9D6E4"})
        set_cell_shading(cover.cell(r, 0), PALE_BLUE)
        text_run(cover.cell(r, 0).paragraphs[0], labels[r], 10, bold=True, color=NAVY)
        text_run(cover.cell(r, 1).paragraphs[0], values[r], 10.5, color=TEXT)
    doc.add_paragraph()
    add_note(doc, "Purpose.", "This guide explains the delivered shop system, the team responsibilities around it, and the operating habits that keep sales, stock and profit records reliable.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_run(p, "A practical guide for daily shop operations and financial control.", 11, italic=True, color=MUTED)
    doc.add_page_break()

    heading(doc, "1. Project overview")
    para(doc, "ADONAK Electronics Shop is a complete retail management system for selling electronics, managing customers and staff, recording payments, issuing invoices, controlling inventory and reviewing financial performance.")
    add_table(doc, ["Area", "What the system supports"], [
        ("Shopfront", "Products, customer accounts, cart, checkout and order history."),
        ("Operations", "Stock, suppliers, product pricing, order handling, delivery and customer records."),
        ("Accountability", "Payment references, invoices, served-by records, staff attendance and audit trails."),
        ("Finance", "Sales analytics, gross profit, operating expenses and payroll cost tracking."),
    ], [2500, 6620])

    heading(doc, "2. What has been delivered")
    for item in [
        "A customer-facing shop for browsing products, managing a cart and placing orders.",
        "An administration area for products, stock, orders, customers, invoices, payments, staff and reports.",
        "Role-based staff access, so staff only see the work areas assigned to them.",
        "Payment and invoice records that identify the handling staff member for accountability.",
        "Lipa Pole Pole support, including the down payment, balance tracking and payment method on the invoice.",
        "Financial reporting that separates settled sales, VAT, product cost, gross profit, payroll and operating expenses.",
    ]:
        add_bullet(doc, item)
    add_note(doc, "Important.", "The project keeps an auditable business record: sensitive actions such as key updates, payment handling, payroll and expense changes are recorded against the user who performed them.")

    heading(doc, "3. Access and user roles")
    para(doc, "Access is role-based. A staff member’s menus and quick actions are shown only when the matching permission has been granted. This reduces mistakes and keeps duties clearly separated.")
    add_table(doc, ["User group", "Main access"], [
        ("Administrators", "Full shop administration, finance reports, staff management, setup and oversight."),
        ("Staff members", "Only assigned operational areas, such as orders, stock, customers, invoices or payment lookup."),
        ("Customers", "Their own shop account, cart, orders, wallet activity and eligible reviews."),
    ], [2500, 6620])
    para(doc, "Client responsibility: create each team member’s own account, use the correct role, and remove or deactivate access immediately when someone leaves or changes duties.", bold_lead="Client responsibility:")

    heading(doc, "4. Core operational modules")
    heading(doc, "Product catalogue and inventory", 2)
    para(doc, "The product catalogue holds each item’s name, category, brand, price, buying cost, stock quantity, supplier details and product image. The warehouse and stock tools help the team monitor availability and identify low-stock items.")
    add_bullet(doc, "Record the buying price whenever a product is added or its supplier cost changes.")
    add_bullet(doc, "Keep the selling price and available stock up to date before listing an item for customers.")
    add_bullet(doc, "Use the low-stock view to plan reorders before products run out.")

    heading(doc, "Orders, payments and delivery", 2)
    para(doc, "Orders can be followed from creation through payment and delivery. Payment references can be searched in the system, while cancellations and refunds are excluded from recognised settled revenue.")
    add_bullet(doc, "Use the correct order status as work progresses; do not mark an order delivered until the goods have reached the customer.")
    add_bullet(doc, "For mobile-money evidence, keep the payment reference accurate. The lookup searches saved shop records; it is not a live verification service from the mobile-money provider.")

    heading(doc, "Lipa Pole Pole", 2)
    para(doc, "The Lipa Pole Pole workflow supports a staged purchase. A customer makes the required wallet down payment, the system records the plan and remaining balance, and each invoice shows the payment method used. This gives the business a clear record of how the customer is paying.")
    add_bullet(doc, "Issue a provisional invoice while a balance remains outstanding.")
    add_bullet(doc, "Issue the final tax receipt only after the order is fully paid and has not been cancelled or refunded.")

    heading(doc, "Invoices and tax records", 2)
    para(doc, "The invoice archive provides a traceable record for completed sales, Lipa Pole Pole payments and cancelled transactions. Invoices identify the staff member who served or handled the transaction, supporting customer service follow-up and internal accountability.")
    add_bullet(doc, "Use the invoice search by order or invoice number when a customer needs a copy.")
    add_bullet(doc, "Do not treat an invoice as a final tax receipt if the order still has an outstanding balance.")

    heading(doc, "Customers and reviews", 2)
    para(doc, "Customer records and order history can be viewed by the appropriate team members. Reviews are controlled so that customers can only submit a product review after a delivered purchase; staff can moderate reviews before they appear publicly.")

    heading(doc, "Staff attendance", 2)
    para(doc, "The staff dashboard supports clock-in and clock-out records, helping management see attendance and shift activity. Attendance information is available to authorised staff and administrators.")
    doc.add_page_break()

    heading(doc, "5. Financial control and profit reporting")
    para(doc, "The Financial Analytics page turns recorded activity into a clear financial view for a selected period. It recognises only fully paid, non-cancelled and non-refunded orders as settled revenue.")
    add_table(doc, ["Measure", "How it is calculated"], [
        ("Gross settled sales", "The full value of qualifying, fully settled orders."),
        ("Net settled sales", "Gross settled sales after the stored VAT component."),
        ("Gross product profit", "Net settled sales less the buying cost saved with each sold item."),
        ("Net profit", "Gross product profit less active operating expenses and paid payroll cost."),
    ], [2900, 6220])
    add_note(doc, "Why buying cost matters.", "When a sale is completed, the system keeps the product cost for that sale. New product costs therefore affect future sales, while historic sales keep their recorded cost for reliable reporting.")

    heading(doc, "Operating expenses", 2)
    para(doc, "The expense module records business costs that are not product purchases, including transport, rent, utilities, airtime or internet, delivery, repairs, marketing, licences, office supplies, security, insurance and bank charges.")
    add_bullet(doc, "Record each expense with its date, category, amount, description, payment method and reference.")
    add_bullet(doc, "Use voiding instead of deleting an incorrect expense, and record the reason. Voided expenses are not included in net profit.")

    heading(doc, "Payroll", 2)
    para(doc, "Payroll supports salary profiles and monthly payroll runs. A payroll run moves from draft to paid, with payment date, method, reference and processing staff member captured for each paid run.")
    add_bullet(doc, "Set each employee’s monthly basic salary and any regular allowances before creating a payroll run.")
    add_bullet(doc, "The business salary expense is the gross pay: basic salary plus allowances. Employee deductions reduce the amount paid to the employee but do not reduce the business’s salary cost.")
    add_bullet(doc, "Void a paid payroll only when necessary, using a clear reason; voided payroll is excluded from net profit.")

    heading(doc, "Historical buying-cost completion", 2)
    para(doc, "Historical sales were reviewed so that profit reporting can use buying costs instead of assuming every sale is pure profit. Where an old sale did not have a saved cost, the then-current product buying cost was recorded as an estimate and an audit record was retained. All new sales now capture their buying cost automatically at checkout.")

    heading(doc, "6. Daily and monthly operating routine")
    add_table(doc, ["When", "Recommended routine"], [
        ("Every sale", "Confirm stock and price, process the correct payment flow, update the order status, and issue the appropriate invoice or receipt."),
        ("Every stock purchase", "Update the product buying price, selling price where required, stock quantity and supplier details."),
        ("Every business expense", "Record it on the expense page on the same day, including the payment reference and method."),
        ("At delivery", "Confirm the goods have reached the customer before marking the order delivered."),
        ("Month end", "Create and pay the monthly payroll, then review Financial Analytics for the period."),
        ("Weekly", "Review low stock, pending orders, customer queries, staff attendance and backup status."),
    ], [2200, 6920])

    heading(doc, "7. Recommended operating procedures")
    heading(doc, "Adding a new product", 2)
    for item in [
        "Open product management and add the product name, category, brand and image.",
        "Enter the buying price, selling price, opening stock and supplier details.",
        "Check the product on the shopfront before making it available to customers.",
    ]:
        add_number(doc, item)
    heading(doc, "Closing a monthly financial period", 2)
    for item in [
        "Ensure all paid orders have the correct status and that refunds or cancellations are recorded accurately.",
        "Enter all operating expenses for the month and complete the payroll run.",
        "Open Financial Analytics, choose the period and review revenue, VAT, product cost, gross profit, expenses, payroll and net profit.",
        "Save or print the report together with supporting invoices and payment references for your business records.",
    ]:
        add_number(doc, item)

    heading(doc, "8. Security, records and good practice")
    for item in [
        "Use individual logins; do not share administrator passwords or staff accounts.",
        "Give staff the minimum access needed for their job and review permissions regularly.",
        "Use documented cancellation, refund and void processes rather than removing records.",
        "Keep payment references, supplier documents and expense evidence safely for audit and reconciliation.",
        "Maintain regular backups and keep a copy outside the main computer where possible.",
    ]:
        add_bullet(doc, item)
    add_note(doc, "Record integrity.", "The system is designed to keep a history of key business actions. Accurate entries at the time of work are the best way to maintain trustworthy stock, customer and profit reports.")

    heading(doc, "9. Handover checklist")
    add_table(doc, ["Item", "Handover status"], [
        ("Customer shop and ordering workflow", "Delivered"),
        ("Product, stock and supplier management", "Delivered"),
        ("Payment records, invoices and Lipa Pole Pole tracking", "Delivered"),
        ("Role-based staff access and attendance", "Delivered"),
        ("Sales, VAT, profit, expense and payroll reporting", "Delivered"),
        ("Historical cost completion for profit reporting", "Completed"),
    ], [6100, 3020])
    para(doc, "This document should be retained with the business records and shared with the staff members responsible for administration, sales, stock and finance.")

    doc.core_properties.title = "ADONAK Electronics Shop — Client Project Documentation"
    doc.core_properties.subject = "Client handover guide"
    doc.core_properties.author = "ADONAK Electronics Shop"
    doc.core_properties.comments = "Client-facing project documentation and operating guide."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
